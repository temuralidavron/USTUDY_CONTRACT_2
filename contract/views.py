import os
import qrcode
import base64
from io import BytesIO

from django.core.files.base import ContentFile
from django.http import HttpResponseBadRequest
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from datetime import timedelta
from dateutil.relativedelta import relativedelta
from django.utils.timezone import localtime

from account.utils import admin_required
from config import settings
from .forms import ContractForm, AdminContractForm, ContractAdminForm
from .models import Contract

from contract.models import Contract
from .utils import link_callback


def home(request):
    unconfirmed_count = 0
    if request.user.is_authenticated and request.user.is_admin:
        unconfirmed_count = Contract.objects.filter(saved=False).count()
    return render(request, 'home.html', {'unconfirmed_count': unconfirmed_count})

# contract/views.py
from django.contrib.auth.decorators import login_required
from .forms import ContractForm
from .models import Contract

@login_required
def my_contracts(request):
    try:
        contract = Contract.objects.filter(user=request.user).first()
    except Contract.DoesNotExist:
        contract = None

    if request.method == "POST":
        if contract:
            form = ContractForm(request.POST, instance=contract)
        else:
            form = ContractForm(request.POST)

        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user
            instance.save()
            return render(request, "contract/contract.html", {"data": contract})
    else:
        form = ContractForm(instance=contract)

    return render(request, "contract/contract.html", {"form": form, "contract": contract})


from django.urls import reverse

@login_required
def create_contract(request):
    # hidden_fields = ['document_given_by', 'document_given_date']
    if request.method == "POST":
        form = ContractForm(request.POST, request.FILES)
        if form.is_valid():
            contract = form.save(commit=False)
            contract.user = request.user

            signature_data = form.cleaned_data.get('signature_data')
            if signature_data:
                format, imgstr = signature_data.split(';base64,')
                ext = format.split('/')[-1]
                contract.signature = ContentFile(base64.b64decode(imgstr), name=f"signature_{request.user.id}.{ext}")

            contract.save()
            return redirect('contract:contract_detail', contract_id=contract.id)
    else:
        form = ContractForm()
    hidden_fields = [
        "parent_full_name",
        "parent_document_series",
        "parent_jshshir",
        "parent_document_given_by",
        "parent_document_given_date",
    ]
    return render(request, "contract/contract_form.html", {"form": form, "hidden_fields": hidden_fields})


@login_required
def contract_detail(request, contract_id):
    contract = get_object_or_404(Contract, id=contract_id, user=request.user)
    return render(request, "contract/contract.html", {"data": contract})

#
# @login_required
# def create_contract(request):
#     if Contract.objects.filter(user=request.user).exists():
#         return redirect('contract:my_contracts')
#     hidden_fields = ['document_series', 'jshshir', 'document_given_by', 'document_given_date']
#
#     if request.method == "POST":
#         form = ContractForm(request.POST)
#         if form.is_valid():
#             contract = form.save(commit=False)
#             contract.user = request.user
#             contract.save()
#             return render(request, "contract/contract.html", {"data": contract})
#     else:
#         form = ContractForm()
#     return render(request, "contract/contract_form.html", {"form": form,"hidden_fields": hidden_fields})



from django.views.decorators.http import require_POST
from django.shortcuts import get_object_or_404
@require_POST
@login_required
def confirm_contract(request):
    contract_id = request.POST.get("contract_id")
    if not contract_id or not contract_id.isdigit():
        return HttpResponseBadRequest("Noto'g'ri yoki yo'q contract_id")

    contract = get_object_or_404(Contract, id=contract_id, user=request.user)
    contract.is_confirmed = True
    contract.save()
    return redirect('contract:home')




@admin_required
def unconfirmed_contracts(request):
    contracts = Contract.objects.filter(saved=False).order_by("-created_at")
    return render(request, 'contract/unconfirmed_contracts.html', {"contracts": contracts})

# contract/forms.py

@admin_required
def admin_fill_contract(request, pk):
    contract = get_object_or_404(Contract, pk=pk, is_confirmed=False)

    if request.method == "POST":
        form = AdminContractForm(request.POST, instance=contract)
        if form.is_valid():
            form.save()
            contract.saved = True
            contract.save()

            # PDF/Word generation qismi bu yerda bo'lishi mumkin (keyinroq qo‘shamiz)

            return redirect('contract:unconfirmed_contracts')
    else:
        form = AdminContractForm(instance=contract)

    return render(request, "contract/admin_fill_contract.html", {"form": form, "contract": contract})

@admin_required
def update_contract(request, pk):
    contract = get_object_or_404(Contract, pk=pk)

    if request.method == 'POST':
        print(request.POST)
        form = ContractAdminForm(request.POST, instance=contract)
        if form.is_valid():
            print(form.cleaned_data)
            form.save()
            return redirect('contract:confirmed_contracts')
        else:
            print(form.errors)
    else:
        form = ContractAdminForm(instance=contract)

    return render(request, 'contract/update_contract.html', {
        'form': form,
        'data': contract    })


@admin_required
def admin_preview_contract(request, pk):
    contract = get_object_or_404(Contract, pk=pk)
    return render(request, "contract/contract.html", {"contract": contract, "admin_preview": True})



@admin_required
@require_POST
def admin_finalize_contract(request, pk):
    contract = get_object_or_404(Contract, pk=pk)

    # ✅ Word fayl yaratish
    from docx import Document
    import io

    doc = Document()
    doc.add_heading("Shartnoma", 0)
    doc.add_paragraph(f"F.I.SH: {contract.full_name}")
    doc.add_paragraph(f"Yosh: {contract.age}")
    doc.add_paragraph(f"Manzil: {contract.address}")
    doc.add_paragraph(f"Kurs: {contract.course_type}")
    doc.add_paragraph(f"1-oylik narx: {contract.initial_price} so‘m")
    doc.add_paragraph(f"Har oylik narx: {contract.price} so‘m")
    doc.add_paragraph(f"Kurs muddati: {contract.monthly_duration} so‘m")
    # boshqa maydonlar...

    # Faylga saqlash yoki yuborish
    word_io = io.BytesIO()
    doc.save(word_io)

    # ✅ PDF yaratish
    from xhtml2pdf import pisa
    from django.template.loader import render_to_string

    html = render_to_string("contract/pdf_template.html", {"contract": contract})
    pdf_io = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html), dest=pdf_io)

    # ✅ Statusni yangilash
    contract.is_confirmed = True
    contract.saved = True
    contract.save()

    return redirect('contract:confirmed_contracts')



# views.py
@admin_required
def confirmed_contracts(request):
    contracts = Contract.objects.filter(saved=True).order_by("-created_at")
    return render(request, 'contract/confirmed_contracts.html', {"contracts": contracts})


from django.template.loader import get_template
from xhtml2pdf import pisa
from django.http import HttpResponse
from .models import Contract
import io

# @admin_required
def download_contract_pdf(request, pk):
    contract = get_object_or_404(Contract, pk=pk)
    template = get_template("contract/save_contract.html")
    start_date = localtime(contract.created_at).date()
    duration = contract.monthly_duration
    if duration is None:
        duration = 1
    sum_course = (contract.price * (duration-1))+contract.initial_price


    intervals = [start_date + relativedelta(months=i) for i in range(int(duration))]
    # QR
    # qr_url = request.build_absolute_uri(
    #     reverse("http://185.8.213.46:8081/", kwargs={"contract_id": contract.id})
    # )
    qr_url = f"http://185.8.213.46:8081/contracts/{contract.id}/"

    # 🔹 QR kod yaratish
    qr = qrcode.make(qr_url)
    buffer = BytesIO()
    qr.save(buffer, format="PNG")
    qr_base64 = base64.b64encode(buffer.getvalue()).decode()

    html = template.render({
        "data": contract,
        "admin_preview": True,
        "intervals": intervals,
        'sum_course': sum_course,
        "show_download_links": False,
        "qr_code": qr_base64,

    })





    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename=contract_{contract.id}.pdf'

    pisa_status = pisa.CreatePDF(html, dest=response,link_callback=link_callback)
    if pisa_status.err:
        return HttpResponse('Xatolik yuz berdi PDF yaratishda')
    return response


@admin_required
def download_contract_word(request, pk):
    contract = get_object_or_404(Contract, pk=pk)
    template = get_template("contract/save_contract.html")
    start_date = localtime(contract.created_at).date()
    duration = contract.monthly_duration
    if duration is None:
        duration = 1
    sum_course = (contract.price * (duration - 1)) + contract.initial_price
    print(sum_course)

    intervals = [start_date + relativedelta(months=i) for i in range(int(duration))]
    html = template.render({
        "data": contract,
        "admin_preview": True,
        "intervals": intervals,
        'sum_course': sum_course,
        "show_download_links": False
    })


    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = f'attachment; filename=contract_{contract.id}.doc'
    response.write(html)
    return response
